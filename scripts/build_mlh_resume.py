"""Build Shaunak Rane's ATS-safe MLH Fellowship resume."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Shaunak_Rane_MLH_Resume.docx"
PDF_OUTPUT = ROOT / "Shaunak_Rane_MLH_Resume.pdf"

# compact_reference_guide with a named "one_page_resume" override.
TOKENS = {
    "font": "Arial",
    "ink": RGBColor(22, 31, 45),
    "muted": RGBColor(75, 85, 99),
    "accent": RGBColor(17, 79, 107),
    "rule": "CBD5E1",
    "body_size": 9.4,
    "body_line": 1.0,
    "section_size": 10.4,
    "page_top": 0.43,
    "page_bottom": 0.42,
    "page_left": 0.58,
    "page_right": 0.58,
}


def set_run_font(
    run,
    *,
    size=None,
    bold=None,
    italic=None,
    color=None,
    underline=None,
):
    run.font.name = TOKENS["font"]
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), TOKENS["font"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if underline is not None:
        run.underline = underline


def set_paragraph_spacing(
    paragraph,
    *,
    before=0,
    after=0,
    line=1.0,
    keep_with_next=False,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def add_hyperlink(paragraph, text, url, *, size=8.5, color=None, bold=False):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), TOKENS["font"])
    fonts.set(qn("w:hAnsi"), TOKENS["font"])
    run_properties.append(fonts)

    run_color = OxmlElement("w:color")
    rgb = color or TOKENS["accent"]
    run_color.set(qn("w:val"), str(rgb))
    run_properties.append(run_color)

    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(run_size)

    if bold:
        run_properties.append(OxmlElement("w:b"))

    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def set_bottom_border(paragraph, *, color=None, size="6", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        p_borders = OxmlElement("w:pBdr")
        p_pr.append(p_borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color or TOKENS["rule"])
    p_borders.append(bottom)


def set_right_tab(paragraph, position_inches=7.34):
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(position_inches),
        WD_TAB_ALIGNMENT.RIGHT,
    )


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Resume Section")
    set_paragraph_spacing(
        paragraph,
        before=4.7,
        after=2.9,
        line=1.0,
        keep_with_next=True,
    )
    run = paragraph.add_run(text.upper())
    set_run_font(
        run,
        size=TOKENS["section_size"],
        bold=True,
        color=TOKENS["accent"],
    )
    set_bottom_border(paragraph)
    return paragraph


def add_entry_heading(doc, left, right, *, after=0.8):
    paragraph = doc.add_paragraph(style="Resume Entry")
    set_right_tab(paragraph)
    set_paragraph_spacing(
        paragraph,
        before=0.4,
        after=after,
        line=1.0,
        keep_with_next=True,
    )
    left_run = paragraph.add_run(left)
    set_run_font(left_run, size=9.65, bold=True, color=TOKENS["ink"])
    paragraph.add_run("\t")
    right_run = paragraph.add_run(right)
    set_run_font(right_run, size=8.8, bold=True, color=TOKENS["muted"])
    return paragraph


def add_project_heading(doc, title, stack, url):
    paragraph = doc.add_paragraph(style="Resume Entry")
    set_right_tab(paragraph)
    set_paragraph_spacing(
        paragraph,
        before=0.5,
        after=0.7,
        line=1.0,
        keep_with_next=True,
    )
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=9.55, bold=True, color=TOKENS["ink"])
    stack_run = paragraph.add_run(f" | {stack}")
    set_run_font(stack_run, size=8.55, color=TOKENS["muted"])
    paragraph.add_run("\t")
    add_hyperlink(paragraph, "GitHub", url, size=8.45, bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="Resume Bullet")
    set_paragraph_spacing(paragraph, before=0, after=1.4, line=1.05)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.13)
    run = paragraph.add_run(text)
    set_run_font(run, size=TOKENS["body_size"], color=TOKENS["ink"])
    return paragraph


def add_skills_line(doc, label, text):
    paragraph = doc.add_paragraph(style="Resume Body")
    set_paragraph_spacing(paragraph, before=0, after=1.4, line=1.02)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.2, bold=True, color=TOKENS["ink"])
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=9.2, color=TOKENS["ink"])
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = TOKENS["ink"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = TOKENS["body_line"]

    for style_name in ("Resume Section", "Resume Entry", "Resume Body"):
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        style.base_style = normal
        style.font.name = TOKENS["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])

    if "Resume Bullet" not in doc.styles:
        bullet = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = doc.styles["Resume Bullet"]
    bullet.base_style = doc.styles["List Bullet"]
    bullet.font.name = TOKENS["font"]
    bullet._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    bullet.font.size = Pt(TOKENS["body_size"])


def build_resume():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(TOKENS["page_top"])
    section.bottom_margin = Inches(TOKENS["page_bottom"])
    section.left_margin = Inches(TOKENS["page_left"])
    section.right_margin = Inches(TOKENS["page_right"])
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    configure_styles(doc)

    properties = doc.core_properties
    properties.title = "Shaunak Rane - MLH Fellowship Resume"
    properties.subject = "AI/ML and Software Engineering Resume"
    properties.author = "Shaunak Rane"
    properties.keywords = (
        "Python, machine learning, PyTorch Geometric, FastAPI, graph neural networks, "
        "time-series validation, software engineering"
    )

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, before=0, after=0.4, line=1.0, keep_with_next=True)
    name_run = name.add_run("SHAUNAK RANE")
    set_run_font(name_run, size=21.5, bold=True, color=TOKENS["ink"])

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(headline, before=0, after=1.2, line=1.0, keep_with_next=True)
    headline_run = headline.add_run(
        "AI/ML Undergraduate | Data Science, Machine Learning & Backend Engineering"
    )
    set_run_font(headline_run, size=9.5, bold=True, color=TOKENS["accent"])

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, before=0, after=0.5, line=1.0, keep_with_next=True)
    add_hyperlink(
        contact,
        "+91 93202 21211",
        "tel:+919320221211",
        size=8.45,
        color=TOKENS["muted"],
    )
    separator = contact.add_run("  |  ")
    set_run_font(separator, size=8.45, color=TOKENS["muted"])
    add_hyperlink(
        contact,
        "shaunakrane914@gmail.com",
        "mailto:shaunakrane914@gmail.com",
        size=8.45,
        color=TOKENS["muted"],
    )
    separator = contact.add_run("  |  ")
    set_run_font(separator, size=8.45, color=TOKENS["muted"])
    location = contact.add_run("Dombivli, Maharashtra")
    set_run_font(location, size=8.45, color=TOKENS["muted"])

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(links, before=0, after=2.3, line=1.0, keep_with_next=True)
    add_hyperlink(
        links,
        "github.com/Shaunakrane914",
        "https://github.com/Shaunakrane914",
        size=8.35,
    )
    separator = links.add_run("  |  ")
    set_run_font(separator, size=8.35, color=TOKENS["muted"])
    add_hyperlink(
        links,
        "sr914.netlify.app",
        "https://sr914.netlify.app",
        size=8.35,
    )

    add_section_heading(doc, "Profile")
    profile = doc.add_paragraph(style="Resume Body")
    set_paragraph_spacing(profile, before=0, after=1.4, line=1.0)
    profile_run = profile.add_run(
        "AI/ML undergraduate with two Univitt Technologies internships spanning software "
        "engineering, data science, and machine learning. Experienced with Python, PyTorch "
        "Geometric, FastAPI, Docker, causal validation, and production-minded testing."
    )
    set_run_font(profile_run, size=9.15, color=TOKENS["ink"])

    add_section_heading(doc, "Education")
    add_entry_heading(
        doc,
        "Universal AI University | B.Tech, Artificial Intelligence & Machine Learning",
        "Expected 2028",
    )
    education = doc.add_paragraph(style="Resume Body")
    set_paragraph_spacing(education, before=0, after=0.8, line=1.0)
    education_run = education.add_run("Karjat, Maharashtra")
    set_run_font(education_run, size=8.85, color=TOKENS["muted"])

    add_section_heading(doc, "Experience")
    add_entry_heading(
        doc,
        "AI Intern | Univitt Technologies",
        "Jul 2026 - Present",
    )
    add_bullet(
        doc,
        "Developed a physics-grounded condition-monitoring platform for centrifugal "
        "compressors across 3 plants and more than 4,000 daily historian records.",
    )
    add_bullet(
        doc,
        "Reproduced 3,879 historical degradation-index (DFI) rows to <1e-10 absolute "
        "difference and added causal maintenance-event logic with 17/17 regression "
        "checks passing.",
    )
    add_bullet(
        doc,
        "Audited 400+ process and mechanical variables, evaluated statistical and ML "
        "forecasts, and used Docker for reproducible backend workflows; retained only "
        "outputs that passed promotion gates.",
    )

    add_entry_heading(
        doc,
        "Software Intern | Univitt Technologies",
        "Apr 2025 - Jul 2025",
        after=0.7,
    )
    add_bullet(
        doc,
        "Built a food-optimization and cost-management program for a Sodexo canteen "
        "using FastAPI and SQL.",
    )
    add_bullet(
        doc,
        "Implemented menu generation, bill-of-materials calculations, inventory and cost "
        "workflows, meal-attendance forecasting, and operator-facing views.",
    )

    add_section_heading(doc, "Projects")
    add_project_heading(
        doc,
        "TopoFlow GNN",
        "Python, PyTorch Geometric, FastAPI",
        "https://github.com/Shaunakrane914/Flow",
    )
    add_bullet(
        doc,
        "Benchmarked GraphSAGE against Kozeny-Carman physics across 1,231 pore-network "
        "samples from 5 formations, using pore-size heterogeneity to select the modeling regime.",
    )
    add_bullet(
        doc,
        "Reduced MSE by 46.2% on Savonnieres and 28.4% on Estaillades while retaining "
        "physics baselines for the other 3 formations.",
    )
    add_bullet(
        doc,
        "Implemented pore-network processing, model training, FastAPI inference, and "
        "scientific visualizations with PyTorch Geometric.",
    )

    add_project_heading(
        doc,
        "Project Aegis",
        "FastAPI, Gemini API, Supabase, WebSockets",
        "https://github.com/Shaunakrane914/Misinformation",
    )
    add_bullet(
        doc,
        "Built backend orchestration for a threat-intelligence prototype with "
        "schema-constrained Gemini outputs, API-key isolation, fallback routing, "
        "REST observability, and WebSocket telemetry.",
    )

    add_project_heading(
        doc,
        "Gridium Protocol",
        "Python, FastAPI, React Three Fiber, Web3",
        "https://github.com/Shaunakrane914/Live-Ai-1",
    )
    add_bullet(
        doc,
        "In a 3-person team, implemented the Python AI engine, backend integration, "
        "and 3D visualization for a 15-node simulated microgrid with DDPG control.",
    )

    add_section_heading(doc, "Engineering Practice")
    add_bullet(
        doc,
        "Reproducibility: versioned experiments, Docker workflows, regression checks, "
        "causal rolling-origin evaluation, and documented promotion gates.",
    )
    add_bullet(
        doc,
        "Collaboration: defined API boundaries across AI, backend, and 3D components "
        "in a 3-person team and maintained setup, architecture, and research documentation.",
    )

    add_section_heading(doc, "Technical Skills")
    add_skills_line(
        doc,
        "Languages",
        "Python, TypeScript/JavaScript, C, SQL",
    )
    add_skills_line(
        doc,
        "ML & Data",
        "PyTorch, PyTorch Geometric, scikit-learn, pandas, NumPy, statsmodels",
    )
    add_skills_line(
        doc,
        "Backend & Web",
        "FastAPI, Flask, React, REST APIs, WebSockets, PostgreSQL, MySQL, Supabase",
    )
    add_skills_line(
        doc,
        "Engineering",
        "Docker, Git/GitHub, pytest, API validation, causal time-series evaluation, data visualization",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


def register_pdf_fonts():
    font_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("Arial", str(font_dir / "arial.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Bold", str(font_dir / "arialbd.ttf")))
    pdfmetrics.registerFont(TTFont("Arial-Italic", str(font_dir / "ariali.ttf")))


def pdf_paragraph(text, style, *, link=None):
    if link:
        text = f'<a href="{link}" color="#114F6B">{text}</a>'
    return Paragraph(text, style)


def build_pdf():
    register_pdf_fonts()
    ink = colors.HexColor("#161F2D")
    muted = colors.HexColor("#4B5563")
    accent = colors.HexColor("#114F6B")
    rule = colors.HexColor("#CBD5E1")

    styles = {
        "name": ParagraphStyle(
            "ResumeName",
            fontName="Arial-Bold",
            fontSize=22.5,
            leading=23,
            textColor=ink,
            alignment=TA_CENTER,
            spaceAfter=1,
        ),
        "headline": ParagraphStyle(
            "ResumeHeadline",
            fontName="Arial-Bold",
            fontSize=9.7,
            leading=10.6,
            textColor=accent,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            fontName="Arial",
            fontSize=8.6,
            leading=9.45,
            textColor=muted,
            alignment=TA_CENTER,
            spaceAfter=1,
        ),
        "section": ParagraphStyle(
            "ResumeSection",
            fontName="Arial-Bold",
            fontSize=10.95,
            leading=11.5,
            textColor=accent,
            alignment=TA_LEFT,
            spaceBefore=5.0,
            spaceAfter=0.5,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            fontName="Arial",
            fontSize=9.7,
            leading=11.55,
            textColor=ink,
            alignment=TA_LEFT,
            spaceAfter=1.0,
        ),
        "entry": ParagraphStyle(
            "ResumeEntry",
            fontName="Arial-Bold",
            fontSize=9.7,
            leading=10.85,
            textColor=ink,
            alignment=TA_LEFT,
            spaceBefore=0.3,
            spaceAfter=0.8,
        ),
        "project": ParagraphStyle(
            "ResumeProject",
            fontName="Arial",
            fontSize=9.35,
            leading=10.7,
            textColor=ink,
            alignment=TA_LEFT,
            spaceBefore=0.3,
            spaceAfter=0.7,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            fontName="Arial",
            fontSize=9.55,
            leading=11.25,
            textColor=ink,
            alignment=TA_LEFT,
            leftIndent=13,
            firstLineIndent=-8,
            bulletIndent=0,
            spaceAfter=1.05,
        ),
        "skills": ParagraphStyle(
            "ResumeSkills",
            fontName="Arial",
            fontSize=9.35,
            leading=10.65,
            textColor=ink,
            alignment=TA_LEFT,
            spaceAfter=0.8,
        ),
    }

    pdf = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        rightMargin=TOKENS["page_right"] * inch,
        leftMargin=TOKENS["page_left"] * inch,
        topMargin=TOKENS["page_top"] * inch,
        bottomMargin=TOKENS["page_bottom"] * inch,
        title="Shaunak Rane - MLH Fellowship Resume",
        author="Shaunak Rane",
        subject="AI/ML and Software Engineering Resume",
    )

    story = []

    def section(title):
        story.append(Paragraph(title.upper(), styles["section"]))
        story.append(
            HRFlowable(
                width="100%",
                thickness=0.55,
                color=rule,
                spaceBefore=0,
                spaceAfter=2.5,
            )
        )

    def bullet(text):
        story.append(Paragraph(text, styles["bullet"], bulletText="-"))

    story.append(Paragraph("SHAUNAK RANE", styles["name"]))
    story.append(
        Paragraph(
            "AI/ML Undergraduate | Data Science, Machine Learning &amp; Backend Engineering",
            styles["headline"],
        )
    )
    story.append(
        Paragraph(
            '<a href="tel:+919320221211" color="#4B5563">+91 93202 21211</a>'
            '  |  <a href="mailto:shaunakrane914@gmail.com" color="#4B5563">'
            "shaunakrane914@gmail.com</a>  |  Dombivli, Maharashtra",
            styles["contact"],
        )
    )
    story.append(
        Paragraph(
            '<a href="https://github.com/Shaunakrane914" color="#114F6B">'
            "github.com/Shaunakrane914</a>  |  "
            '<a href="https://sr914.netlify.app" color="#114F6B">'
            "sr914.netlify.app</a>",
            styles["contact"],
        )
    )

    section("Profile")
    story.append(
        Paragraph(
            "AI/ML undergraduate with two Univitt Technologies internships spanning "
            "software engineering, data science, and machine learning. Experienced "
            "with Python, PyTorch Geometric, FastAPI, Docker, causal validation, and "
            "production-minded testing.",
            styles["body"],
        )
    )

    section("Education")
    story.append(
        Paragraph(
            "<b>Universal AI University | B.Tech, Artificial Intelligence &amp; "
            "Machine Learning</b> | Expected 2028",
            styles["entry"],
        )
    )
    story.append(Paragraph("Karjat, Maharashtra", styles["body"]))

    section("Experience")
    story.append(
        Paragraph(
            "<b>AI Intern | Univitt Technologies</b> | "
            "Jul 2026 - Present",
            styles["entry"],
        )
    )
    bullet(
        "Developed a physics-grounded condition-monitoring platform for centrifugal "
        "compressors across 3 plants and more than 4,000 daily historian records."
    )
    bullet(
        "Reproduced 3,879 historical degradation-index (DFI) rows to &lt;1e-10 "
        "absolute difference and added causal maintenance-event logic with 17/17 "
        "regression checks passing."
    )
    bullet(
        "Audited 400+ process and mechanical variables, evaluated statistical and ML "
        "forecasts, and used Docker for reproducible backend workflows; retained only "
        "outputs that passed promotion gates."
    )
    story.append(
        Paragraph(
            "<b>Software Intern | Univitt Technologies</b> | "
            "Apr 2025 - Jul 2025",
            styles["entry"],
        )
    )
    bullet(
        "Built a food-optimization and cost-management program for a Sodexo canteen "
        "using FastAPI and SQL."
    )
    bullet(
        "Implemented menu generation, bill-of-materials calculations, inventory and "
        "cost workflows, meal-attendance forecasting, and operator-facing views."
    )

    section("Projects")
    story.append(
        Paragraph(
            '<b>TopoFlow GNN</b> | Python, PyTorch Geometric, FastAPI | '
            '<a href="https://github.com/Shaunakrane914/Flow" color="#114F6B">'
            "GitHub</a>",
            styles["project"],
        )
    )
    bullet(
        "Benchmarked GraphSAGE against Kozeny-Carman physics across 1,231 "
        "pore-network samples from 5 formations, using pore-size heterogeneity "
        "to select the modeling regime."
    )
    bullet(
        "Reduced MSE by 46.2% on Savonnieres and 28.4% on Estaillades while "
        "retaining physics baselines for the other 3 formations."
    )
    bullet(
        "Implemented pore-network processing, model training, FastAPI inference, "
        "and scientific visualizations with PyTorch Geometric."
    )
    story.append(
        Paragraph(
            '<b>Project Aegis</b> | FastAPI, Gemini API, Supabase, WebSockets | '
            '<a href="https://github.com/Shaunakrane914/Misinformation" '
            'color="#114F6B">GitHub</a>',
            styles["project"],
        )
    )
    bullet(
        "Built backend orchestration for a threat-intelligence prototype with "
        "schema-constrained Gemini outputs, API-key isolation, fallback routing, "
        "REST observability, and WebSocket telemetry."
    )
    story.append(
        Paragraph(
            '<b>Gridium Protocol</b> | Python, FastAPI, React Three Fiber, Web3 | '
            '<a href="https://github.com/Shaunakrane914/Live-Ai-1" color="#114F6B">'
            "GitHub</a>",
            styles["project"],
        )
    )
    bullet(
        "In a 3-person team, implemented the Python AI engine, backend integration, "
        "and 3D visualization for a 15-node simulated microgrid with DDPG control."
    )

    section("Engineering Practice")
    bullet(
        "Reproducibility: versioned experiments, Docker workflows, regression checks, "
        "causal rolling-origin evaluation, and documented promotion gates."
    )
    bullet(
        "Collaboration: defined API boundaries across AI, backend, and 3D components "
        "in a 3-person team and maintained setup, architecture, and research documentation."
    )

    section("Technical Skills")
    story.append(
        Paragraph(
            "<b>Languages:</b> Python, TypeScript/JavaScript, C, SQL",
            styles["skills"],
        )
    )
    story.append(
        Paragraph(
            "<b>ML &amp; Data:</b> PyTorch, PyTorch Geometric, scikit-learn, "
            "pandas, NumPy, statsmodels",
            styles["skills"],
        )
    )
    story.append(
        Paragraph(
            "<b>Backend &amp; Web:</b> FastAPI, Flask, React, REST APIs, "
            "WebSockets, PostgreSQL, MySQL, Supabase",
            styles["skills"],
        )
    )
    story.append(
        Paragraph(
            "<b>Engineering:</b> Docker, Git/GitHub, pytest, API validation, causal "
            "time-series evaluation, data visualization",
            styles["skills"],
        )
    )

    pdf.build(story)
    print(PDF_OUTPUT)


if __name__ == "__main__":
    build_resume()
    build_pdf()
